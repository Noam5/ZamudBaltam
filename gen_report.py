# -*- coding: utf-8 -*-
# gen_report.py — build the Hebrew RTL .docx Baltam report from report_data.json
import json, os, sys
try: sys.stdout.reconfigure(encoding='utf-8')
except Exception: pass
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

D = json.load(open(os.path.join(os.path.dirname(__file__), 'report_data.json'), encoding='utf-8'))
L = D['levers']
mo = lambda x: f"{round(x/4):,}"          # ₪/month from 4-month total
GREEN = RGBColor(0x21, 0x73, 0x46)
ORANGE = RGBColor(0xC6, 0x51, 0x00)

doc = Document()
# base style
st = doc.styles['Normal']
st.font.name = 'Arial'; st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn('w:cs'), 'Arial')

def rtl_p(p):
    pPr = p._p.get_or_add_pPr()
    b = OxmlElement('w:bidi'); pPr.append(b)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return p

def para(text='', bold=False, size=11, color=None, space_after=6):
    p = doc.add_paragraph(); rtl_p(p)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        r = p.add_run(text); r.bold = bold; r.font.size = Pt(size)
        if color: r.font.color.rgb = color
    return p

def heading(text, level=1):
    sizes = {1: 16, 2: 13}
    p = doc.add_paragraph(); rtl_p(p)
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(sizes.get(level, 12))
    r.font.color.rgb = GREEN if level == 1 else ORANGE
    return p

def bullets(items):
    for it in items:
        p = doc.add_paragraph(style='List Bullet'); rtl_p(p)
        p.paragraph_format.space_after = Pt(2)
        if isinstance(it, tuple):
            r = p.add_run(it[0]); r.bold = True
            p.add_run(it[1])
        else:
            p.add_run(it)

def shade(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), color); tcPr.append(shd)

def add_table(headers, rows, header_fill='217346', widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.RIGHT
    # RTL column order
    t._tbl.tblPr.append(OxmlElement('w:bidiVisual'))
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        shade(hdr[i], header_fill)
        p = hdr[i].paragraphs[0]; rtl_p(p)
        r = p.add_run(str(h)); r.bold = True; r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); r.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            p = cells[i].paragraphs[0]; rtl_p(p)
            r = p.add_run(str(v)); r.font.size = Pt(10)
            if i == 0: r.bold = True
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

# ---- make whole section RTL ----
sectPr = doc.sections[0]._sectPr
sectPr.append(OxmlElement('w:bidi'))

# ===================== TITLE =====================
tp = doc.add_paragraph(); rtl_p(tp); tp.paragraph_format.space_after = Pt(2)
tr = tp.add_run('דו״ח ייעול עלויות רכב ממוגן (בלת״ם)'); tr.bold = True; tr.font.size = Pt(20); tr.font.color.rgb = GREEN
sp = para('ניתוח 312 נסיעות אמיתיות, פברואר–מאי 2026 · נכון ל-06/06/2026', size=10);
sp.runs[0].font.color.rgb = RGBColor(0x66,0x66,0x66)

# ===================== EXEC SUMMARY =====================
heading('תקציר מנהלים — השורה התחתונה')
para('כיום משולם לזמוד סכום קבוע של ₪52,000 לחודש עבור הרכב הצמוד, ובנוסף מחויב בלת״ם '
     'פר-נסיעה כאשר הרכב הצמוד תפוס. מטרת הדו״ח: לצמצם את הבלת״ם.', space_after=8)
bullets([
    ('בלת״ם נוכחי: ', f'כ-₪{mo(L["baseGreedy"])} לחודש (₪{L["baseGreedy"]:,} ל-4 חודשים).'),
    ('בלי שום עלות נוספת ', f'(שיבוץ חכם + גמישות פגישות מציאותית + מיזוג נסיעות): ירידה לכ-₪{mo(L["realPool"])} לחודש — חיסכון של כ-₪{mo(L["baseGreedy"]-L["realPool"])} לחודש (כ-36%).'),
    ('עם רכב משני חלקי ', f'(כ-7 שעות שבועיות בחלונות החמים בלבד): ניתן לרדת לכיוון ₪1,000–1,500 לחודש — משתלם כל עוד הרכב המשני עולה פחות מ-₪3,500 לחודש.'),
    ('הסיבה השורשית: ', '‎100% מהתנגשויות הבלת״ם הן שני רופאים שונים שזקוקים לרכב הממוגן היחיד באותו הזמן — לא בזבוז שיבוץ. לכן המנוף האמיתי הוא קיבולת ופיזור ביקוש, לא תחכום בלוח הזמנים.'),
])

# ===================== LEVER MENU =====================
heading('תפריט המנופים — כמה כל אחד שווה')
para('כל הסכומים לחודש. כל שלב מצטבר על קודמו.', size=10, space_after=4)
add_table(
    ['מנוף', 'בלת״ם נותר (₪/ח׳)', 'חיסכון מצטבר/ח׳', 'עלות', 'שורה תחתונה'],
    [
        ['מצב נוכחי', f'₪{mo(L["baseGreedy"])}', '—', '—', 'בסיס'],
        ['שיבוץ חכם של הצמוד', f'₪{mo(L["optimalDispatch"])}', f'₪{mo(L["baseGreedy"]-L["optimalDispatch"])}', 'חינם', 'ליישם מיד'],
        ['+ גמישות פגישות (מציאותי)', f'₪{mo(L["realNoPool"])}', f'₪{mo(L["baseGreedy"]-L["realNoPool"])}', 'תיאום מרפאות', 'מומלץ'],
        ['+ מיזוג נסיעות (2 רופאים)', f'₪{mo(L["realPool"])}', f'₪{mo(L["baseGreedy"]-L["realPool"])}', 'תיאום', 'מנוף קטן'],
        ['רכב משני חלקי', '~₪1,000–1,500', 'עד ~₪4,850', 'מחיר הרכב', 'תלוי מחיר'],
        ['הרחבת שעות הצמוד', f'מקסימום ₪{mo(488*4)}', '—', 'עלות לשעה', 'כמעט לא משתלם'],
    ],
    widths=[5.2, 3.2, 3.2, 2.8, 2.6]
)

# ===================== CORE FINDING =====================
heading('הממצא המרכזי')
para('כל נסיעת בלת״ם מסוג "התנגשות" (B) נובעת משני רופאים שונים שזקוקים לרכב הממוגן '
     'היחיד בו-זמנית — אף פעם לא מהתנגשות בין נסיעות של אותו רופא. גם שיבוץ אופטימלי של '
     f'רכב אחד מוריד את הבלת״ם רק לכ-₪{mo(L["optimalDispatch"])} לחודש: הביקוש פשוט עולה על רכב אחד בשעות השיא. '
     'מכאן ששלוש דרכי הפעולה האמיתיות הן: (א) לפזר ביקוש (הזזת פגישות), (ב) להגדיל קיבולת '
     '(רכב משני חלקי / מיזוג נוסעים), (ג) לוודא שהצמוד אכן מבצע כל נסיעה שביכולתו.', space_after=8)

# ===================== HOTSPOTS =====================
heading('מוקדי ההתנגשות החוזרים (לפי יום ומסלול)')
para('אלו הדפוסים השבועיים החוזרים שבהם נוצר בלת״ם — היעדים לפיזור/קיבולת. סכומים ל-4 חודשים.', size=10)
add_table(['יום ומסלול', 'מס׳ פעמים', '₪ (4 ח׳)', '₪/חודש'],
          [[r['k'], r['n'], f"₪{r['sum']:,}", f"₪{r['perMo']:,}"] for r in D['byRoute']],
          widths=[8.0, 2.2, 2.6, 2.4])

heading('בלת״ם לפי רופא ויום', level=2)
add_table(['יום ורופא', 'מס׳ פעמים', '₪/חודש'],
          [[r['k'], r['n'], f"₪{round(r['sum']/4):,}"] for r in D['byDoctorDay']],
          widths=[7.0, 2.6, 2.6])

# ===================== 2ND CAR =====================
heading('רכב משני חלקי — היכן ובכמה')
para('רכב ממוגן שני נדרש רק בחלונות צרים וקבועים — לא במשרה מלאה. החלונות המומלצים '
     '(לאחר שיבוץ חכם + גמישות מציאותית):', space_after=4)
add_table(['יום', 'חלון', 'נסיעות', '₪/חודש'],
          [[w['day'], f"{w['start']}–{w['end']}", w['count'], f"₪{w['perMo']:,}"] for w in D['windows'] if w['perMo'] > 0],
          widths=[1.6, 3.2, 2.0, 2.4])
para('כלל החלטה: רכב משני חלקי בחלונות אלו (כ-7 שעות שבועיות) משתלם אם עלותו נמוכה '
     'מ-כ-₪3,500 לחודש (כ-₪115 לשעה). כדאי לבקש הצעת מחיר הממוקדת בחלונות אלו בלבד.', space_after=8)

# ===================== POOLING =====================
heading('מיזוג נסיעות (2 רופאים ברכב אחד)')
para(f'הרכב הממוגן הוא בן 7 מקומות, ושיתוף נסיעות מותר. אולם בנתונים אלו המיזוג הוא '
     f'המנוף הקטן ביותר — כ-₪{mo(L["realNoPool"]-L["realPool"])} לחודש בלבד. הסיבה: ההוצאה הגדולה ביותר '
     '(נסיעות נטליה לאלון מורה בימי שני) היא יעד מרוחק שאליו רק היא נוסעת — אין עם מי לאחד. '
     'הדפוס המשתלם היחיד והחוזר: בימי שישי לירון (אלפי מנשה↔עמנואל) ועבד (אלפי מנשה↔קדומים) '
     'נוסעים באותן שעות — רכב אחד יכול לבצע אלפי מנשה→עמנואל→קדומים ולהוריד את שניהם.', space_after=8)

# ===================== HOURS =====================
heading('הרחבת שעות הצמוד — למה לא')
para('חלון השעות הנוכחי כבר טוב (כבר הורחב בימי שני/רביעי). יתרת הפוטנציאל קטנה '
     f'(כ-₪{mo(488*4)} לחודש בלבד), וההרחבה כרוכה בעלות לשעה. ההחזר לשעה נוסף נמוך '
     '(רביעי כ-₪73 לשעה, שלישי כ-₪167 לשעה) — לרוב לא משתלם. אין לרדוף אחרי מנוף זה.', space_after=8)

# ===================== ACTION PLAN =====================
heading('תוכנית פעולה מתועדפת')
bullets([
    ('1. שיבוץ חכם (חינם, מיידי): ', 'לוודא שכל נסיעה שהרכב הצמוד יכול לבצע אכן מבוצעת על-ידו ולא מועברת לאורציון. חיסכון כ-₪1,025 לחודש.'),
    ('2. גמישות פגישות (תיאום מרפאות): ', 'להזיז פגישות ב-15–30 דק׳ במוקדים החוזרים (שישי בוקר, ראשון). חיסכון נוסף כ-₪770 לחודש.'),
    ('3. מיזוג בשישי: ', 'לאחד את לירון ועבד לרכב אחד לעמנואל/קדומים. כ-₪300 לחודש.'),
    ('4. רכב משני חלקי: ', 'לבקש הצעת מחיר לחלונות החמים בלבד (טבלת הרכב המשני). להפעיל אם מתחת ל-₪3,500 לחודש.'),
    ('5. לא להרחיב שעות: ', 'העלות לשעה גבוהה מהחיסכון.'),
])

# ===================== APPENDIX =====================
heading('נספח — שיטה, הנחות ואזהרות')
bullets([
    ('בסיס נתונים: ', f'{D["counts"]["trips"]} נסיעות אמיתיות מחשבוניות אורציון (פברואר–מאי 2026); מתוכן {D["counts"]["B"]} התנגשות (B) ו-{D["counts"]["BML"]} מחוץ לשעות (BML).'),
    ('מנוע החישוב: ', 'מנוע הסיווג של האפליקציה (Z/B/BML), בעזרת טבלת המחירים והשעות החיות. כל המספרים נבדקו אוטומטית מול ערכי-עוגן.'),
    ('הנחת מודל החיוב: ', '₪52,000 קבוע לצמוד + חיוב פר-נסיעה לבלת״ם מעליו (אושר על-ידך).'),
    ('מספרי הגמישות הם אומדן: ', 'מבוססים על מודל שיבוץ; שיבוץ חכם ורכב משני מדויקים יותר.'),
    ('מסלולים חסרים: ', 'מעט נסיעות (בי״ח מאיר, צור יצחק נחל צלמון) חסרות בטבלת המחירים ולכן אינן מתומחרות — השפעתן זניחה.'),
    ('כלי אינטראקטיבי: ', 'לשונית "ייעול בלת״ם" באפליקציה מאפשרת לשחק עם כל המנופים ולראות את ההשפעה בזמן אמת.'),
])

out = os.path.join(os.path.dirname(__file__), 'דוח_ייעול_בלתם.docx')
doc.save(out)
print('saved report (', os.path.getsize(out), 'bytes )')
